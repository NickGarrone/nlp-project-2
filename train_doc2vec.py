import os
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import string
from gensim.models.doc2vec import Doc2Vec, TaggedDocument


# Text Extraction Code
def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_pdf_file(file_path):
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
    return text

def read_word_file(file_path):
    doc = docx.Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def read_files_in_directory(directory_path):
    texts = {}
    
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        
        if filename.endswith('.txt'):
            texts[filename] = read_text_file(file_path)
        elif filename.endswith('.pdf'):
            texts[filename] = read_pdf_file(file_path)
        elif filename.endswith('.docx'):
            texts[filename] = read_word_file(file_path)
    return texts


# Text Preprocessing Code
# Download NLTK resources
nltk.download('punkt')
nltk.download('stopwords')

# Preprocessing function
def preprocess_text(text):
    # Lowercase
    text = text.lower()
    
    # Tokenization
    tokens = word_tokenize(text)
    
    # Remove punctuation and stop words
    tokens = [word for word in tokens if word.isalnum()]  # Remove punctuation
    tokens = [word for word in tokens if word not in stopwords.words('english')]  # Remove stop words
    
    return tokens

# Prepare tagged documents for Doc2Vec
def prepare_tagged_documents(files_contents):
    tagged_documents = []
    
    for idx, (filename, content) in enumerate(files_contents.items()):
        tokens = preprocess_text(content)
        tagged_documents.append(TaggedDocument(tokens, [f'doc_{idx}']))
    
    return tagged_documents


# Example usage
directory_path = 'training_data'
files_contents = read_files_in_directory(directory_path)
tagged_data = prepare_tagged_documents(files_contents)

# Now you can train your Doc2Vec model
model = Doc2Vec(vector_size=20, min_count=1, epochs=100)

# Build vocabulary
model.build_vocab(tagged_data)

# Train the model
model.train(tagged_data, total_examples=model.corpus_count, epochs=model.epochs)

# Save the model for future use
model.save("doc2vec_model.d2v")


# similar_docs = model.dv.most_similar('doc_0')
# print(similar_docs)
